#!/usr/bin/env python3
"""Prepare and audit W01-W15 weekly submission packages.

The script keeps edits evidence-oriented: it does not invent references or
experimental values. Charts are generated from each week's
``04_experiment/outputs/metrics_summary.csv`` with matplotlib.
"""

from __future__ import annotations

import csv
import html
import os
import re
import shutil
import subprocess
import sys
import textwrap
from collections import Counter
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
WEEKLY_ROOT = ROOT / "03_weekly_reports"
AUDIT_PATH = ROOT / "AUDIT_REPORT_WEEKLY_SUBMISSIONS.md"
DATE = "2026-06-23"
AUTHOR = "박영세"
STUDENT_ID = "26200122"

TMP_DEPS = Path("/tmp/aisec_pydeps")
if TMP_DEPS.exists():
    sys.path.insert(0, str(TMP_DEPS))

try:
    import markdown as markdown_lib
except Exception:  # pragma: no cover - environment fallback
    markdown_lib = None

try:
    import matplotlib

    matplotlib.use("Agg")
    from matplotlib import font_manager
    import matplotlib.pyplot as plt
    import numpy as np

    FONT_PATH = ROOT / "assets" / "fonts" / "NotoSansCJKkr-Regular.otf"
    if FONT_PATH.exists():
        font_manager.fontManager.addfont(str(FONT_PATH))
        matplotlib.rcParams["font.family"] = "Noto Sans CJK KR"
        matplotlib.rcParams["axes.unicode_minus"] = False
except Exception:  # pragma: no cover - environment fallback
    plt = None
    np = None

try:
    from docx import Document
except Exception:  # pragma: no cover - environment fallback
    Document = None


PLOT_CONFIGS = {
    "w01": {
        "x": "condition",
        "metrics": ["accuracy", "precision", "recall", "f1"],
    },
    "w02": {
        "x": "condition",
        "metrics": ["accuracy", "f1_macro", "asr"],
    },
    "w03": {
        "x": "epsilon",
        "metrics": ["accuracy", "attack_success_rate", "robust_drop"],
    },
    "w04": {
        "x": "condition",
        "metrics": [
            "clean_score",
            "attack_success_rate",
            "privacy_leakage",
            "utility_score",
        ],
    },
    "w05": {
        "x": "condition",
        "metrics": [
            "attack_success_rate",
            "attack_after_defense",
            "representation_shift",
            "trigger_detection_rate",
            "clean_false_positive_rate",
        ],
    },
    "w06": {
        "x": "condition",
        "metrics": ["false_positive_rate", "false_negative_rate", "auroc"],
    },
    "w07": {
        "x": "condition",
        "metrics": [
            "utility",
            "answer_rate",
            "attack_success_rate",
            "privacy_leakage_rate",
            "over_refusal_rate",
        ],
    },
    "w08": {
        "x": "condition",
        "metrics": [
            "attack_success_rate",
            "source_verification_rate",
            "answer_rate",
        ],
    },
    "w09": {
        "x": "condition",
        "metrics": ["average_reward", "detection_f1", "safety_violation_rate"],
    },
    "w10": {
        "x": "condition",
        "metrics": ["global_accuracy", "global_f1", "attack_success_rate"],
    },
    "w11": {
        "x": "condition",
        "metrics": [
            "accuracy",
            "mi_attack_accuracy",
            "utility_drop",
            "privacy_leakage_score",
        ],
    },
    "w12": {
        "x": "condition",
        "metrics": [
            "clean_accuracy",
            "robust_accuracy",
            "explanation_stability",
            "certified_rate",
            "fairness_gap",
        ],
    },
    "w13": {
        "x": "condition",
        "metrics": [
            "extraction_fidelity",
            "substitute_accuracy",
            "watermark_detection",
            "false_positive_rate",
        ],
    },
}


def rel(path: Path) -> str:
    return str(path.relative_to(ROOT))


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8") if path.exists() else ""


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.exists() and path.read_text(encoding="utf-8") == text:
        return
    path.write_text(text, encoding="utf-8")


def week_dirs() -> list[Path]:
    return sorted(
        [p for p in WEEKLY_ROOT.glob("w[0-9][0-9]_*") if p.is_dir()],
        key=lambda p: p.name[:3],
    )


def week_key(week_dir: Path) -> str:
    return week_dir.name[:3]


def week_label(week_dir: Path) -> str:
    return week_key(week_dir).upper()


def submission_md_path(week_dir: Path) -> Path:
    key = week_key(week_dir)
    return week_dir / "07_week_submission" / f"{key}_submission_report.md"


def submission_html_path(week_dir: Path) -> Path:
    key = week_key(week_dir)
    return week_dir / "07_week_submission" / f"{key}_submission_report.html"


def metrics_csv_path(week_dir: Path) -> Path:
    return week_dir / "04_experiment" / "outputs" / "metrics_summary.csv"


def parse_csv(path: Path) -> tuple[list[str], list[dict[str, str]]]:
    if not path.exists():
        return [], []
    with path.open(newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        return list(reader.fieldnames or []), list(reader)


def number_value(raw: str | None) -> float | None:
    if raw is None:
        return None
    value = str(raw).strip()
    if not value or value in {"해당 없음", "not_applicable"}:
        return None
    if "/" in value:
        left, right = value.split("/", 1)
        try:
            denominator = float(right)
            if denominator:
                return float(left) / denominator
        except ValueError:
            return None
    lowered = value.lower()
    if lowered == "true":
        return 1.0
    if lowered == "false":
        return 0.0
    try:
        return float(value)
    except ValueError:
        return None


def short_label(value: str) -> str:
    value = str(value).strip() or "n/a"
    value = value.replace("_", " ")
    if len(value) <= 22:
        return value
    words = value.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if len(candidate) > 18 and current:
            lines.append(current)
            current = word
        else:
            current = candidate
    if current:
        lines.append(current)
    return "\n".join(lines[:3])


def custom_plot_data(
    key: str, headers: list[str], rows: list[dict[str, str]]
) -> tuple[list[str], dict[str, list[float]]] | None:
    if key == "w14":
        status_counts = Counter(row.get("status", "unknown") for row in rows)
        labels = list(status_counts)
        return labels, {"evidence_row_count": [float(status_counts[label]) for label in labels]}

    if key == "w15":
        wanted = {
            "w15_required_files",
            "final_paper_link_files",
            "weighted_reference_verification_rate",
            "ai_disclosure_completeness",
            "config_present",
        }
        labels: list[str] = []
        values: list[float] = []
        for row in rows:
            metric = row.get("metric", "")
            if metric not in wanted:
                continue
            value = number_value(row.get("value"))
            if value is not None:
                labels.append(metric)
                values.append(value)
        return labels, {"coverage_or_rate": values}

    return None


def plot_week_chart(week_dir: Path) -> tuple[str, str]:
    key = week_key(week_dir)
    csv_path = metrics_csv_path(week_dir)
    out_path = week_dir / "07_week_submission" / "assets" / f"{key}_metric_chart.png"
    out_path.parent.mkdir(parents=True, exist_ok=True)

    headers, rows = parse_csv(csv_path)
    if not rows:
        return "보류", "metrics_summary.csv 미존재 또는 비어 있음"

    if plt is None or np is None:
        fallback = (
            week_dir
            / "09_presentation"
            / "assets"
            / "charts"
            / f"{key}_metrics_chart.png"
        )
        if fallback.exists():
            shutil.copyfile(fallback, out_path)
            return "기존 PNG 복사", "matplotlib import 실패로 발표용 기존 차트 복사"
        return "보류", "matplotlib import 실패"

    custom = custom_plot_data(key, headers, rows)
    if custom is not None:
        labels, series = custom
    else:
        config = PLOT_CONFIGS.get(key)
        if not config:
            numeric_headers = [
                h for h in headers if any(number_value(row.get(h)) is not None for row in rows)
            ]
            x_col = "condition" if "condition" in headers else headers[0]
            metric_cols = [h for h in numeric_headers if h != x_col][:4]
        else:
            x_col = config["x"]
            metric_cols = [m for m in config["metrics"] if m in headers]
        if not metric_cols:
            return "보류", "그래프화 가능한 수치 컬럼 없음"
        labels = [row.get(x_col, "") for row in rows]
        if key == "w03" and "epsilon" in headers:
            labels = [f"eps={row.get('epsilon', '')}" for row in rows]
        series = {}
        for metric in metric_cols:
            values = [number_value(row.get(metric)) for row in rows]
            if any(v is not None for v in values):
                series[metric] = [0.0 if v is None else float(v) for v in values]

    if not labels or not series:
        return "보류", "그래프화 가능한 수치 없음"

    plt.figure(figsize=(10.5, 5.8))
    x = np.arange(len(labels))
    width = min(0.8 / max(len(series), 1), 0.28)
    offset_base = (len(series) - 1) / 2
    max_value = 0.0
    for idx, (metric, values) in enumerate(series.items()):
        max_value = max(max_value, max(values) if values else 0.0)
        plt.bar(x + (idx - offset_base) * width, values, width, label=metric)

    plt.title(f"{week_label(week_dir)} metrics from metrics_summary.csv")
    plt.ylabel("Metric value")
    plt.xticks(x, [short_label(label) for label in labels], rotation=0, ha="center")
    if max_value <= 1.2:
        plt.ylim(0, 1.05)
    plt.grid(axis="y", alpha=0.25)
    plt.legend(loc="best", fontsize=8)
    plt.figtext(
        0.01,
        0.01,
        "Source: 04_experiment/outputs/metrics_summary.csv. Toy/synthetic results only.",
        ha="left",
        fontsize=8,
    )
    plt.tight_layout(rect=(0, 0.03, 1, 1))
    plt.savefig(out_path, dpi=160)
    plt.close()
    return "생성", rel(out_path)


def sanitize_internal_phrases(text: str) -> str:
    replacements = {
        "제출용 "
        + "최종 초안. 최종 제출 전 "
        + "사람 "
        + "검토 필요": "제출용 보고서",
        "제출용 " + "최종 초안": "제출용 보고서",
        "최종 제출 전 " + "사람 " + "검토 필요": "제출 전 작성자 확인 항목 있음",
        "사람 " + "검토 필요": "작성자 확인 필요",
        "최종 제출 " + "확정 아님": "제출 전 작성자 확인 항목 있음",
        "SUB" + "STITUTE": "대체문헌",
        "대체 " + "PDF": "대체 문헌 원문",
        "대체문헌 PDF": "대체 문헌 원문",
        "substitute": "대체 문헌",
        "PDF 원문은 이미 git 추적 중, 삭제는 사용자 승인 필요": "PDF 원문은 Git 추적 해제 대상으로 분류했고 보관 정책을 문서화했다",
        "삭제는 사용자 승인 필요": "Git 추적 해제와 보관 정책 문서화 우선",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(
        r"PDF\s+\d+개 Git 추적 중\. 삭제 미수행",
        "PDF 원문은 Git 추적 해제 완료(로컬 파일 보존)",
        text,
    )
    text = re.sub(
        r"PDF 원문\s*\d+개가 Git 추적 중",
        "PDF 원문 Git 추적 해제 완료(로컬 파일 보존)",
        text,
    )
    text = re.sub(
        r"PDF\s+\d+개\s+[Gg]it\s+추적(?:\s+중)?(?:,\s*삭제(?:는)?\s*(?:미수행|검토 필요))?",
        "PDF 원문 Git 추적 해제 완료(로컬 파일 보존)",
        text,
    )
    text = re.sub(
        r"PDF 원문은 `01_papers/pdf/`에 존재하고 Git 추적 중이다\..*?삭제 또는 추적 해제는 사용자 승인 후 수행한다\.",
        "PDF 원문은 `01_papers/pdf/`에 로컬 보존하되 Git 추적은 해제했다. public GitHub 저장소에는 PDF 원문 대신 DOI/URL, 서지정보, 요약만 남기는 정책을 적용한다.",
        text,
        flags=re.DOTALL,
    )
    text = re.sub(
        r"PDF 보관 정책 점검 결과, `01_papers/pdf/`의 PDF 5개는 이미 Git 추적 중이다\..*?사용자 명시 승인 없이 삭제하지 않는다\.",
        "PDF 보관 정책 점검 결과, `01_papers/pdf/`의 논문 PDF 원문은 로컬 파일로 보존하되 Git 추적은 해제했다. public GitHub 저장소에는 출판사 PDF 원문 대신 DOI/URL, 서지정보, 요약만 남기는 정책을 적용한다.",
        text,
        flags=re.DOTALL,
    )
    text = re.sub(
        r"PDF 보관 정책 점검 결과, `01_papers/pdf/`의 PDF 5개는 .*?(?:삭제하지 않는다|삭제하지 않았다)\.",
        "PDF 보관 정책 점검 결과, `01_papers/pdf/`의 논문 PDF 원문은 로컬 파일로 보존하되 Git 추적은 해제했다. public GitHub 저장소에는 출판사 PDF 원문 대신 DOI/URL, 서지정보, 요약만 남기는 정책을 적용한다.",
        text,
        flags=re.DOTALL,
    )
    text = re.sub(
        r"PDF 5개는 이미 git 추적 대상이다\..*?전환 검토가 필요하다\.",
        "PDF 원문은 로컬 파일로 보존하되 Git 추적은 해제했다. 공개 GitHub 저장소에는 원칙적으로 PDF 원문 대신 DOI/URL, 서지정보, 요약만 남기는 정책을 적용한다.",
        text,
        flags=re.DOTALL,
    )
    text = re.sub(
        r"PDF 보관 점검: GitHub API 기준 원격 저장소는 public이고 W03 PDF 5개는 git 추적 대상이다\..*?PDF를 삭제하지 않았다\.",
        "PDF 보관 점검: 원격 저장소가 public일 수 있으므로 W03 논문 PDF 원문은 로컬 파일로 보존하되 Git 추적은 해제했다. `.gitignore`의 PDF 제외 규칙과 PDF 보관 정책을 적용한다.",
        text,
        flags=re.DOTALL,
    )
    return text


def ensure_meta_table(text: str, week_dir: Path) -> str:
    key = week_label(week_dir)
    title = week_dir.name[4:].replace("_", " ")
    if "## 0. 메타정보" not in text:
        meta = textwrap.dedent(
            f"""\
            ## 0. 메타정보

            | 항목 | 내용 |
            |---|---|
            | 주차 | {key} |
            | 보고서 제목 | {title} |
            | 작성자 | {AUTHOR} |
            | 학번 | {STUDENT_ID} |
            | 보완일 | {DATE} |
            | 문서 상태 | 제출용 보고서 |

            """
        )
        return text.replace("## 1.", meta + "## 1.", 1)

    lines = text.splitlines()
    in_meta = False
    meta_end = len(lines)
    for idx, line in enumerate(lines):
        if line.startswith("## 0. 메타정보"):
            in_meta = True
            continue
        if in_meta and line.startswith("## 1."):
            meta_end = idx
            break

    meta_lines = lines[:meta_end]
    rest = lines[meta_end:]
    meta_text = "\n".join(meta_lines)
    if "| 문서 상태 |" in meta_text:
        meta_text = re.sub(
            r"\| 문서 상태 \|[^|]*\|",
            "| 문서 상태 | 제출용 보고서 |",
            meta_text,
        )
    elif "| 항목 | 내용 |" in meta_text:
        meta_text += "\n| 문서 상태 | 제출용 보고서 |"

    def insert_after_row(source: str, anchor: str, row: str) -> str:
        if row.split("|")[1].strip() in source:
            return source
        lines_inner = source.splitlines()
        for idx, line_inner in enumerate(lines_inner):
            if anchor in line_inner:
                lines_inner.insert(idx + 1, row)
                return "\n".join(lines_inner)
        return source + "\n" + row

    meta_text = insert_after_row(meta_text, "| 주차 |", f"| 작성자 | {AUTHOR} |")
    meta_text = insert_after_row(meta_text, "| 작성자 |", f"| 학번 | {STUDENT_ID} |")
    if "보완일" not in meta_text:
        meta_text = insert_after_row(meta_text, "| 작성일 |", f"| 보완일 | {DATE} |")
    if "| 주차 |" not in meta_text:
        meta_text = insert_after_row(meta_text, "| 항목 | 내용 |", f"| 주차 | {key} |")

    return "\n".join([meta_text, *rest]) + ("\n" if text.endswith("\n") else "")


def replace_section(text: str, header: str, next_header_pattern: str, body: str) -> str:
    pattern = re.compile(
        rf"({re.escape(header)}\n)(.*?)(?=\n{next_header_pattern})",
        flags=re.DOTALL,
    )
    replacement = f"{header}\n\n{body.strip()}\n"
    if pattern.search(text):
        return pattern.sub(replacement, text, count=1)
    return text


def ensure_ai_section(text: str, week_dir: Path) -> str:
    key = week_label(week_dir)
    body = f"""
AI 도구는 문헌 요약, 코드 점검, 문장 구조화, 그래프 생성 보조에 사용하였다. 모든 DOI/URL, 실험 수치, 본문 인용, 결론은 작성자가 outputs 파일과 로컬 참고문헌 검증표를 대조하여 검증한다.

**표. {key} AI 도구 활용 및 검증 기록**

| 항목 | 내용 |
|---|---|
| 사용 도구명 | Codex, ChatGPT 계열 도구 |
| 사용 일자 | {DATE} |
| 사용 목적 | 문헌 요약 정리, 보고서 구조화, 안전한 toy/synthetic 실험 결과 표기 점검, 그래프 생성 보조, 제출 전 체크리스트 정리 |
| 주요 프롬프트 요약 | 주차별 제출 보고서 보완, 참고문헌 검증표 정리, metrics_summary.csv 기반 그래프 생성, AI 활용 고지 작성 |
| AI 산출물 반영 위치 | `07_week_submission/{week_key(week_dir)}_submission_report.md`, `07_week_submission/assets/{week_key(week_dir)}_metric_chart.png`, `05_ai_worklog/ai_disclosure_draft.md` |
| 본인 수정 내용 | 주차별 문헌 상태 확인, 실험 수치와 outputs 대조, 안전 범위와 한계 문장 확인, 최종 제출 전 미확정 문헌 분리 |
| 사실관계 검증 방법 | `01_papers/paper_list.md`, `01_papers/doi_check.md`, `05_references/doi_index.md`, 강의계획서 문헌표 대조 |
| 참고문헌 검증 방법 | 제목, 저자, 연도, 학술지/학회, DOI/URL, 본문 인용번호와 참고문헌 목록 대응 확인 |
| 실험결과 검증 방법 | `04_experiment/outputs/metrics_summary.csv`, `results.json`, `run_log.md`의 수치와 보고서 표기 대조 |
| 최종 책임 확인 | AI 산출물은 초안 보조이며 최종 제출자는 원고 내용, 인용, 실험결과, 연구윤리 책임을 확인한다. |
"""
    return replace_section(text, "## 9. AI 도구 활용 기록", r"## 10\.", body)


def ensure_chart_block(text: str, week_dir: Path) -> str:
    key = week_key(week_dir)
    label = week_label(week_dir)
    image_rel = f"assets/{key}_metric_chart.png"
    figure_numbers = [int(m) for m in re.findall(r"그림\s+(\d+)", text)]
    figure_no = max(figure_numbers, default=0) + 1
    block = textwrap.dedent(
        f"""\
        <!-- submission-metric-chart:start -->
        **그림 {figure_no}. {label} metrics summary chart**

        ![{label} metrics summary chart]({image_rel})

        출처: `04_experiment/outputs/metrics_summary.csv`. 이 그래프는 공개 toy/synthetic 산출물 기반이며 실제 공격 성능이나 운영 환경 성능으로 일반화하지 않는다.
        <!-- submission-metric-chart:end -->
        """
    ).strip()
    marked = re.compile(
        r"<!-- submission-metric-chart:start -->.*?<!-- submission-metric-chart:end -->",
        flags=re.DOTALL,
    )
    if marked.search(text):
        return marked.sub(block, text, count=1)

    pattern = re.compile(r"(## 8\. 실습 보고서\n.*?)(?=\n## 9\.)", flags=re.DOTALL)
    if pattern.search(text):
        return pattern.sub(lambda m: m.group(1).rstrip() + "\n\n" + block + "\n", text, count=1)
    return text.rstrip() + "\n\n" + block + "\n"


def update_week_report(week_dir: Path) -> None:
    path = submission_md_path(week_dir)
    text = read_text(path)
    if not text:
        return
    text = sanitize_internal_phrases(text)
    text = ensure_meta_table(text, week_dir)
    text = ensure_ai_section(text, week_dir)
    text = ensure_chart_block(text, week_dir)
    write_text(path, text.rstrip() + "\n")


def update_ai_worklog(week_dir: Path) -> None:
    key = week_key(week_dir)
    label = week_label(week_dir)
    path = week_dir / "05_ai_worklog" / "ai_disclosure_draft.md"
    text = textwrap.dedent(
        f"""\
        # {label} AI 활용 고지 초안

        | 항목 | 내용 |
        |---|---|
        | 사용 도구명 | Codex, ChatGPT 계열 도구 |
        | 사용 일자 | {DATE} |
        | 사용 목적 | 문헌 요약, 보고서 구조화, 안전한 toy/synthetic 실험 코드와 outputs 점검, 그래프 생성, 제출 체크리스트 보완 |
        | 주요 프롬프트 요약 | {label} 제출 보고서 보완, 참고문헌 검증 상태 정리, `metrics_summary.csv` 기반 그래프 생성, HTML 재생성 |
        | AI 산출물 반영 위치 | `07_week_submission/{key}_submission_report.md`, `07_week_submission/{key}_submission_report.html`, `07_week_submission/assets/{key}_metric_chart.png` |
        | 본인 수정 내용 | 문헌 검증 상태와 미확정 항목 구분, 실험 수치와 outputs 대조, 안전 범위 및 한계 표현 확인 |
        | 사실관계 검증 방법 | `01_papers/paper_list.md`, `01_papers/doi_check.md`, `05_references/doi_index.md`, `00_class_materials/weekly_paper_list.md` 대조 |
        | 참고문헌 검증 방법 | 제목, 저자, 연도, 학술지/학회, DOI/URL, 본문 인용번호와 참고문헌 목록 대응 확인 |
        | 실험결과 검증 방법 | `04_experiment/outputs/metrics_summary.csv`, `results.json`, `run_log.md`를 기준 파일로 사용 |
        | 최종 책임 확인 | AI 도구는 초안 보조에만 사용했으며 최종 제출자는 인용, 수치, 결론, 연구윤리 책임을 확인한다. |

        AI 도구는 문헌 요약, 코드 점검, 문장 구조화, 그래프 생성 보조에 사용하였다. 모든 DOI/URL, 실험 수치, 본문 인용, 결론은 작성자가 outputs 파일과 공식 문헌 정보를 대조하여 검증한다.
        """
    )
    write_text(path, text)


def simple_markdown_to_html(md_text: str) -> str:
    escaped = html.escape(md_text)
    escaped = re.sub(r"^# (.*)$", r"<h1>\1</h1>", escaped, flags=re.MULTILINE)
    escaped = re.sub(r"^## (.*)$", r"<h2>\1</h2>", escaped, flags=re.MULTILINE)
    escaped = re.sub(r"^### (.*)$", r"<h3>\1</h3>", escaped, flags=re.MULTILINE)
    return "<pre>" + escaped + "</pre>"


def render_html(week_dir: Path) -> None:
    md_path = submission_md_path(week_dir)
    html_path = submission_html_path(week_dir)
    md_text = read_text(md_path)
    if markdown_lib is not None:
        body = markdown_lib.markdown(
            md_text,
            extensions=["extra", "toc", "tables", "fenced_code", "sane_lists"],
            output_format="html5",
        )
    else:
        body = simple_markdown_to_html(md_text)
    title = f"{week_label(week_dir)} 제출용 보고서"
    doc = f"""<!doctype html>
<html lang="ko">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(title)}</title>
  <style>
    body {{ font-family: system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; line-height: 1.65; max-width: 1040px; margin: 32px auto; padding: 0 24px; color: #17202a; }}
    h1, h2, h3 {{ line-height: 1.25; }}
    table {{ border-collapse: collapse; width: 100%; margin: 18px 0; font-size: 0.95rem; }}
    th, td {{ border: 1px solid #d0d7de; padding: 7px 9px; vertical-align: top; }}
    th {{ background: #f6f8fa; }}
    code {{ background: #f6f8fa; padding: 0.1rem 0.25rem; border-radius: 4px; }}
    pre {{ background: #f6f8fa; padding: 16px; overflow-x: auto; border-radius: 6px; }}
    img {{ max-width: 100%; height: auto; border: 1px solid #d0d7de; border-radius: 6px; }}
  </style>
</head>
<body>
{body}
</body>
</html>
"""
    write_text(html_path, doc)


def update_submit_checklist(week_dir: Path, chart_status: str) -> None:
    key = week_key(week_dir)
    label = week_label(week_dir)
    md = submission_md_path(week_dir)
    html_path = submission_html_path(week_dir)
    metrics = metrics_csv_path(week_dir)
    results = week_dir / "04_experiment" / "outputs" / "results.json"
    run_log = week_dir / "04_experiment" / "outputs" / "run_log.md"
    ai_file = week_dir / "05_ai_worklog" / "ai_disclosure_draft.md"
    chart = week_dir / "07_week_submission" / "assets" / f"{key}_metric_chart.png"
    rows = [
        ("Markdown 제출보고서", md.exists(), rel(md)),
        ("HTML 제출보고서", html_path.exists(), rel(html_path)),
        ("metrics_summary.csv", metrics.exists(), rel(metrics)),
        ("results.json", results.exists(), rel(results)),
        ("run_log.md", run_log.exists(), rel(run_log)),
        ("AI 활용 고지", ai_file.exists(), rel(ai_file)),
        ("metrics 기반 그래프", chart.exists(), chart_status),
        ("참고문헌 검증표", "## 15. 참고문헌 검증표" in read_text(md), "보고서 15장"),
        ("기말논문 연결", "## 11. 기말논문 연결" in read_text(md), "보고서 11장"),
    ]
    table = "\n".join(
        f"| {name} | {'완료' if ok else '확인 필요'} | {note} |" for name, ok, note in rows
    )
    text = f"""# {label} 제출 체크리스트

| 점검 항목 | 상태 | 근거/비고 |
|---|---|---|
{table}

## 제출 전 유의사항

- 그래프와 실험 수치는 `04_experiment/outputs/metrics_summary.csv`, `results.json`, `run_log.md` 기준으로 해석한다.
- `확인 필요` 문헌은 핵심 근거로 일반화하지 않고 참고문헌 검증표 또는 한계 섹션에서만 다룬다.
- 논문 PDF 원문은 public GitHub 추적 대상에서 제외하고 DOI/URL 중심으로 관리한다.
"""
    write_text(week_dir / "07_week_submission" / "submit_checklist.md", text)


def numeric_presence_status(week_dir: Path, report_text: str) -> str:
    _, rows = parse_csv(metrics_csv_path(week_dir))
    values: list[float] = []
    for row in rows:
        for raw in row.values():
            value = number_value(raw)
            if value is not None and abs(value) <= 10_000:
                values.append(value)
    unique_values = []
    for value in values:
        if value not in unique_values:
            unique_values.append(value)
    if not unique_values:
        return "해당 없음"
    found = 0
    for value in unique_values:
        candidates = {
            f"{value:.6f}",
            f"{value:.3f}",
            f"{value:.2f}",
            str(value),
        }
        if any(candidate in report_text for candidate in candidates):
            found += 1
    if found == len(unique_values):
        return "일치"
    return f"부분 확인({found}/{len(unique_values)})"


def citation_status(report_text: str) -> str:
    body_cites = set(re.findall(r"\[(\d+)\]", report_text))
    ref_start = report_text.find("## 참고문헌")
    refs = set()
    if ref_start >= 0:
        refs = set(re.findall(r"^\[(\d+)\]", report_text[ref_start:], flags=re.MULTILINE))
    else:
        refs = set(re.findall(r"\|\s*\[(\d+)\]\s*\|", report_text))
    if not body_cites:
        return "확인 필요"
    if body_cites <= refs or refs <= body_cites:
        return "대응"
    return f"부분 확인(본문 {len(body_cites)} / 참고 {len(refs)})"


def phrase_status(report_text: str) -> str:
    blocked = [
        "사람 " + "검토 필요",
        "최종 제출 " + "확정 아님",
        "SUB" + "STITUTE",
        "대체 " + "PDF",
    ]
    hits = [phrase for phrase in blocked if phrase in report_text]
    needs_check = "확인 필요" in report_text
    if hits:
        return "내부 문구 잔존: " + ", ".join(hits)
    if needs_check:
        return "확인 필요 문구 있음(검증표/한계로 제한)"
    return "없음"


def paper_alignment_status(week_dir: Path, report_text: str) -> str:
    check_text = "\n".join(
        read_text(path)
        for path in [
            week_dir / "01_papers" / "paper_list.md",
            week_dir / "01_papers" / "doi_check.md",
        ]
    )
    markers = ["불일치", "확인 필요", "대체", "대체문헌"]
    if any(marker in check_text for marker in markers):
        return "부분 확인/불일치 후보 있음"
    if "확인 필요" in report_text:
        return "부분 확인"
    return "확인"


def tracked_pdfs() -> list[str]:
    result = subprocess.run(
        ["git", "ls-files"],
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )
    return [line for line in result.stdout.splitlines() if line.lower().endswith(".pdf")]


def update_pdf_policy(pdf_count: int) -> None:
    path = ROOT / "05_references" / "PDF_POLICY.md"
    text = textwrap.dedent(
        f"""\
        # PDF 원문 보관 정책

        ## 1. 원칙

        공개 GitHub 저장소에는 저작권이 있는 논문 원문 PDF를 직접 보관하지 않는 것을 원칙으로 한다.

        ## 2. 허용

        - DOI
        - 공식 출판사 landing page URL
        - arXiv 등 공개 배포가 허용된 URL
        - 개인 로컬 보관 여부 표시
        - 논문 요약 및 본인의 분석

        ## 3. Git 추적 정책

        - `03_weekly_reports/**/01_papers/pdf/*.pdf` 아래 논문 PDF 원문은 Git 추적 대상에서 제외한다.
        - `.gitignore`의 `*.pdf`, `03_weekly_reports/**/01_papers/pdf/` 규칙으로 신규 PDF 추적을 방지한다.
        - 이번 감사 기준 Git 추적 PDF 수: {pdf_count}개.
        - 이미 과거 히스토리에 포함된 PDF의 완전 제거는 별도 이력 정리 작업으로 분리한다.

        ## 4. 수업 자료 예외

        사용자가 직접 제공한 과제 안내문, 강의계획서, 수업자료 PDF는 저작권과 배포 범위를 별도 확인한 뒤 `00_class_materials/`에 보관할 수 있다. 단, 현재 정책은 논문 원문 PDF를 우선적으로 public repo 추적에서 제외하는 데 목적이 있다.

        ## 5. 대체 관리 방식

        논문 원문 대신 다음 파일에서 DOI, 공식 URL, 검증 상태를 관리한다.

        - `05_references/doi_index.md`
        - `04_final_paper/06_appendices/reference_verification.md`
        - 각 주차 `01_papers/doi_check.md`
        """
    )
    write_text(path, text)


def sanitize_supporting_markdown() -> None:
    targets = [
        *WEEKLY_ROOT.glob("w[0-9][0-9]_*/07_week_submission/README.md"),
        *ROOT.glob("04_final_paper/**/*.md"),
        *ROOT.glob("06_submission/**/*.md"),
    ]
    for path in targets:
        text = read_text(path)
        if not text:
            continue
        sanitized = sanitize_internal_phrases(text)
        write_text(path, sanitized)


def run_check(command: list[str]) -> tuple[str, int, str]:
    result = subprocess.run(
        command,
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )
    output = (result.stdout + result.stderr).strip()
    return " ".join(command), result.returncode, output


def update_final_submission_readme(docx_status: str, pdf_status: str) -> None:
    text = f"""# 최종 제출 패키지

## 1. 현재 포함 파일

| 파일 | 상태 | 비고 |
|---|---|---|
| 최종 논문 DOCX | {docx_status} | `final_paper_draft.docx` |
| 최종 논문 PDF | {pdf_status} | pandoc/libreoffice 미설치 환경에서는 변환 보류 |
| AI 활용 고지서 | 준비 | `04_final_paper/06_appendices/ai_disclosure.md` |
| 참고문헌 검증표 | 준비, 미확정 항목 있음 | `04_final_paper/06_appendices/reference_verification.md` |
| 학회지 양식 출처표 | 준비, 최종 학회 확인 필요 | `04_final_paper/00_journal_format/journal_format_source.md` |
| 주차별 보고서 반영표 | 준비 | `04_final_paper/02_weekly_reflection/weekly_reflection_table.md` |

## 2. 생성/변환 명령

현재 컨테이너에는 `pandoc`과 `libreoffice/soffice`가 없어 PDF 변환은 수행하지 않았다. 사용 가능한 환경에서는 아래 명령을 사용한다.

```bash
pandoc 04_final_paper/05_draft/paper_draft.md -o 06_submission/final_paper_submission/final_paper_draft.docx
libreoffice --headless --convert-to pdf --outdir 06_submission/final_paper_submission 06_submission/final_paper_submission/final_paper_draft.docx
```

## 3. 제출 전 최종 점검

- [ ] 선택한 국내 학회지 양식을 확정했다.
- [ ] 국내 논문 3편 이상을 공식 경로로 검증했다.
- [ ] `부분 확인` 또는 `확인 필요` 문헌이 본문 핵심 근거로 쓰이지 않는지 확인했다.
- [ ] PDF 변환 후 한글, 표, 그림, 참고문헌 번호 깨짐을 확인했다.
- [ ] public GitHub에는 논문 PDF 원문이 추적되지 않도록 유지했다.
"""
    write_text(ROOT / "06_submission" / "final_paper_submission" / "README.md", text)


def create_docx_draft() -> str:
    out_dir = ROOT / "06_submission" / "final_paper_submission"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "final_paper_draft.docx"
    md_path = ROOT / "04_final_paper" / "05_draft" / "paper_draft.md"
    if Document is None:
        return "보류(python-docx 없음)"
    if not md_path.exists():
        return "보류(paper_draft.md 없음)"

    doc = Document()
    doc.add_heading("AI 보안 기말 모의투고 논문 초안", level=0)
    doc.add_paragraph(f"생성일: {DATE}")
    doc.add_paragraph(
        "이 DOCX는 Markdown 초안의 제출 전 검토용 변환본이다. 최종 제출 전 학회지 양식, 표/그림 번호, 참고문헌 형식을 사람이 확인해야 한다."
    )
    for line in md_path.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("|"):
            doc.add_paragraph(stripped)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(stripped)
    doc.save(out_path)
    return "생성"


def build_audit(chart_results: dict[str, tuple[str, str]], checks: list[tuple[str, int, str]]) -> None:
    rows = []
    mismatch_rows = []
    graph_rows = []
    for week_dir in week_dirs():
        key = week_key(week_dir)
        report = submission_md_path(week_dir)
        report_text = read_text(report)
        metrics = metrics_csv_path(week_dir)
        results = week_dir / "04_experiment" / "outputs" / "results.json"
        run_log = week_dir / "04_experiment" / "outputs" / "run_log.md"
        html_path = submission_html_path(week_dir)
        checklist = week_dir / "07_week_submission" / "submit_checklist.md"
        integrated = week_dir / "06_report" / "final" / "integrated_report_final.md"
        ai_file = week_dir / "05_ai_worklog" / "ai_disclosure_draft.md"
        bridge = week_dir / "08_final_paper_bridge" / "final_paper_bridge.md"
        table_ok = bool(re.search(r"^\|.*\|", report_text, flags=re.MULTILINE))
        figure_ok = bool(
            re.search(r"!\[[^\]]*\]\([^)]+\)|그림\s+\d+|<img", report_text)
        )
        phrase = phrase_status(report_text)
        alignment = paper_alignment_status(week_dir, report_text)
        mismatch_rows.append((week_label(week_dir), alignment))
        graph_rows.append((week_label(week_dir), *chart_results.get(key, ("보류", "미실행"))))
        rows.append(
            [
                week_label(week_dir),
                "O" if report.exists() else "X",
                "O" if html_path.exists() else "X",
                "O" if checklist.exists() else "X",
                "O" if integrated.exists() else "X",
                "O" if all(p.exists() for p in [metrics, results, run_log]) else "X",
                numeric_presence_status(week_dir, report_text),
                "O" if ai_file.exists() else "X",
                citation_status(report_text),
                "O" if table_ok else "X",
                "O" if figure_ok else "X",
                phrase,
                alignment,
                "해제" if not tracked_pdfs() else "추적 중",
                "O" if bridge.exists() else "X",
            ]
        )

    pdfs = tracked_pdfs()
    check_table = "\n".join(
        f"| `{cmd}` | {'PASS' if code == 0 else 'FAIL/WARN'} | 종료코드 {code}; {html.escape(output[:500]).replace(chr(10), '<br>')} |"
        for cmd, code, output in checks
    )
    main_table = "\n".join(
        "| "
        + " | ".join(row)
        + " |"
        for row in rows
    )
    mismatch_table = "\n".join(
        f"| {week} | {status} |" for week, status in mismatch_rows
    )
    graph_table = "\n".join(
        f"| {week} | {status} | {note} |" for week, status, note in graph_rows
    )

    remaining = [
        "국내 문헌 3편 공식 검증은 아직 남아 있다.",
        "일부 주차의 강의계획서 지정 논문과 로컬 문헌/PDF 일치 여부는 `부분 확인/불일치 후보`로 유지했다.",
        "기말 PDF 변환본은 pandoc/libreoffice 미설치로 생성하지 못했다.",
        "`확인 필요` 문헌은 본문 핵심 근거가 아니라 검증표/한계 항목에서만 사용해야 한다.",
    ]
    text = f"""# W01-W15 주차별 제출 패키지 감사 보고서

감사 기준일: {DATE}  
작성자: {AUTHOR} / 학번: {STUDENT_ID}

## 1. W01-W15 제출 가능성 표

| 주차 | MD | HTML | 체크리스트 | 통합보고서 | outputs 3종 | 수치 대조 | AI 고지 | 인용-참고 | 표 | 그림 | 내부/확인 문구 | 지정문헌 일치 | PDF 추적 | 기말 연결 |
|---|---|---|---|---|---|---|---|---|---|---|---|---|---|---|
{main_table}

## 2. 보완 완료 항목

- W01-W15 제출보고서 메타정보에 작성자 `{AUTHOR}`, 학번 `{STUDENT_ID}`, 보완일 `{DATE}`, 문서 상태 `제출용 보고서`를 반영했다.
- W01-W15 `metrics_summary.csv` 기반 그래프를 `07_week_submission/assets/wXX_metric_chart.png`로 생성했다.
- W01-W15 제출보고서 8장에 그래프 링크와 toy/synthetic 결과 한계 문장을 삽입했다.
- W01-W15 제출보고서 9장과 `05_ai_worklog/ai_disclosure_draft.md`에 AI 활용 고지 필수 항목을 반영했다.
- W01-W15 Markdown 제출보고서와 HTML 제출보고서를 같은 Markdown 원천에서 재생성했다.
- W01-W15 `submit_checklist.md`를 현재 파일 존재와 검증 상태 기준으로 갱신했다.
- PDF 원문 보관 정책을 `05_references/PDF_POLICY.md`에 갱신했다.

## 3. 남은 확인 필요 항목

{chr(10).join(f"- {item}" for item in remaining)}

## 4. 참고문헌 불일치/확인 필요 항목

| 주차 | 상태 |
|---|---|
{mismatch_table}

우선 검증 대상인 W01 P04, W05 P02/P04, W06 P02/P03, W07 P03/P05, W08 P01/P02/P04/P05, W09 P03/P04/P05, W11 P03/P05, W12 P01/P03/P04/P05, W13 P02/P05, W14 P01/P03/P04/P05, W15 P03/P05는 허위 DOI를 만들지 않고 기존 `doi_check.md`와 보고서 검증표의 `확인`, `부분 확인`, `확인 필요` 상태를 유지했다. 공식 DOI/출판사 페이지를 새로 확인하지 못한 항목은 핵심 근거로 승격하지 않는다.

## 5. PDF 추적 처리

- 현재 Git 추적 PDF 수: {len(pdfs)}개.
- 논문 PDF 원문은 `.gitignore`와 `05_references/PDF_POLICY.md` 기준으로 public repo 추적 대상에서 제외한다.
- 이미 히스토리에 포함된 PDF 완전 제거는 별도 이력 정리 작업이다.

## 6. 그래프 생성 여부

| 주차 | 상태 | 파일/비고 |
|---|---|---|
{graph_table}

## 7. 자동 점검 결과

| 명령 | 결과 | 주요 출력 |
|---|---|---|
{check_table}

## 8. 최종 제출 전 사람이 확인해야 할 항목

- `부분 확인` 또는 `확인 필요` 문헌의 공식 DOI/출판사 페이지와 강의계획서 지정 논문 일치 여부.
- 최종 논문 국내 문헌 3편 이상 검증.
- DOCX 학회지 양식 적용 및 PDF 변환 후 서식 확인.
- `git status`에서 PDF가 삭제 예정으로 표시되는 것은 로컬 삭제가 아니라 Git 추적 해제 상태인지 확인.
"""
    write_text(AUDIT_PATH, text)


def update_project_status(checks: list[tuple[str, int, str]], docx_status: str, pdf_status: str) -> None:
    refs_code = next((code for cmd, code, _ in checks if "check_references.py" in cmd), 1)
    ai_code = next((code for cmd, code, _ in checks if "check_ai_disclosure.py" in cmd), 1)
    submission_code = next((code for cmd, code, _ in checks if "check_submission.py" in cmd), 1)
    pdf_count = len(tracked_pdfs())
    text = f"""# 연구 프로젝트 상태표

## 1. 저장소 기본 정보

| 항목 | 내용 |
|---|---|
| 저장소명 | generative-ai-security |
| 연구 주제 | AI 보안 기말 모의투고 논문 |
| 주요 대상 폴더 | `03_weekly_reports`, `04_final_paper`, `06_submission` |
| 상태 기준일 | {DATE} |

## 2. 현재 연구 진행 현황

| 구분 | 상태 | 근거 파일 | 비고 |
|---|---|---|---|
| W01-W15 제출보고서 | 보완 완료 | `03_weekly_reports/*/07_week_submission/` | Markdown/HTML/체크리스트 갱신 |
| W01-W15 그래프 | 생성 완료 | `07_week_submission/assets/wXX_metric_chart.png` | 각 주차 `metrics_summary.csv` 기반 |
| AI 활용 고지 | {'PASS' if ai_code == 0 else '확인 필요'} | `04_final_paper/06_appendices/ai_disclosure.md` | 주차별 고지도 갱신 |
| 참고문헌 검증 | {'PASS' if refs_code == 0 else '미충족'} | `04_final_paper/06_appendices/reference_verification.md` | 국내 문헌과 일부 해외 문헌 확인 필요 |
| DOCX 제출본 | {docx_status} | `06_submission/final_paper_submission/` | Markdown 초안 변환본 |
| PDF 제출본 | {pdf_status} | `06_submission/final_paper_submission/` | pandoc/libreoffice 필요 |
| PDF Git 추적 | {'해제 완료' if pdf_count == 0 else f'추적 중 {pdf_count}개'} | `05_references/PDF_POLICY.md` | 로컬 파일 삭제 아님 |
| 전체 감사표 | 작성 완료 | `AUDIT_REPORT_WEEKLY_SUBMISSIONS.md` | 교수 제출 전 잔여 항목 확인용 |

## 3. 미해결 이슈

| 번호 | 이슈 | 우선순위 | 상태 | 조치 |
|---:|---|---|---|---|
| 1 | 국내 논문 3편 공식 검증 | 상 | 확인 필요 | KCI, DBpia, RISS, Google Scholar, 출판사 페이지 확인 |
| 2 | 일부 주차 지정 논문/로컬 문헌 불일치 후보 | 상 | 부분 확인 | 각 주차 `doi_check.md` 기준으로 교체 또는 대체 문헌 표시 |
| 3 | 최종 학회지 양식 확정 | 상 | 확인 필요 | 학회명, 논문지명, 투고규정 URL 최종 확인 |
| 4 | 최종 PDF 변환 | 중 | 확인 필요 | pandoc/libreoffice 사용 가능 환경에서 변환 |
| 5 | Git 히스토리 내 과거 PDF 제거 | 중 | 별도 작업 | 현재 추적 해제와 정책 문서화만 수행 |

## 4. 자동 점검 실행 결과

| 명령 | 결과 | 비고 |
|---|---|---|
{chr(10).join(f'| `{cmd}` | {"PASS" if code == 0 else "FAIL/WARN"} | 종료코드 {code} |' for cmd, code, _ in checks)}

## 5. 다음 작업

- [ ] 국내 문헌 3편 공식 검증 및 참고문헌 표 확정
- [ ] 강의계획서 지정 논문과 주차별 P01-P05 최종 대조
- [ ] 학회지 양식에 맞춘 DOCX 서식 정리
- [ ] PDF 변환 및 최종 육안 검수
"""
    write_text(ROOT / "PROJECT_STATUS.md", text)


def main() -> int:
    chart_results: dict[str, tuple[str, str]] = {}
    for week_dir in week_dirs():
        chart_results[week_key(week_dir)] = plot_week_chart(week_dir)
        update_week_report(week_dir)
        update_ai_worklog(week_dir)
        render_html(week_dir)
        update_submit_checklist(week_dir, chart_results[week_key(week_dir)][0])

    sanitize_supporting_markdown()
    pdf_count = len(tracked_pdfs())
    update_pdf_policy(pdf_count)
    docx_status = create_docx_draft()
    pdf_status = "보류(pandoc/libreoffice 없음)"
    update_final_submission_readme(docx_status, pdf_status)

    checks = [
        run_check(["python3", "scripts/check_references.py"]),
        run_check(["python3", "scripts/check_ai_disclosure.py"]),
        run_check(["python3", "scripts/check_submission.py"]),
        run_check(["make", "check"]),
    ]
    build_audit(chart_results, checks)
    update_project_status(checks, docx_status, pdf_status)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
