from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "paper" / "JKAIS_paper_draft.md"
DOCX_OUT = ROOT / "paper" / "JKAIS_paper_draft.docx"
PDF_OUT = ROOT / "paper" / "JKAIS_paper_draft.pdf"


LINK_RE = re.compile(r"\[([^\]]+)\]\([^)]+\)")
IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")


def clean_inline(text: str) -> str:
    text = IMAGE_RE.sub(r"Figure: \1 (\2)", text)
    text = LINK_RE.sub(r"\1", text)
    text = BOLD_RE.sub(r"\1", text)
    return text.replace("`", "")


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    return bool(stripped.startswith("|") and set(stripped.replace("|", "").replace(":", "").replace("-", "").strip()) == set())


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines: list[str] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        if not is_table_separator(lines[i]):
            cells = [clean_inline(cell.strip()) for cell in lines[i].strip().strip("|").split("|")]
            table_lines.append(cells)
        i += 1
    return table_lines, i


def add_docx_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(width):
            table.cell(r_idx, c_idx).text = row[c_idx] if c_idx < len(row) else ""


def build_docx(lines: list[str]) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Malgun Gothic"
    styles["Normal"].font.size = Pt(10)

    in_code = False
    code_lines: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith("```"):
            if in_code:
                para = doc.add_paragraph()
                run = para.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(8)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.strip().startswith("|"):
            rows, i = parse_table(lines, i)
            add_docx_table(doc, rows)
            continue

        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip("#")), 4)
            doc.add_heading(clean_inline(line.lstrip("#").strip()), level=level)
        elif line.startswith("- "):
            doc.add_paragraph(clean_inline(line[2:].strip()), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(clean_inline(re.sub(r"^\d+\.\s+", "", line)), style="List Number")
        elif line.startswith(">"):
            para = doc.add_paragraph(clean_inline(line.lstrip("> ").strip()))
            para.style = styles["Intense Quote"]
        else:
            doc.add_paragraph(clean_inline(line))
        i += 1

    doc.save(DOCX_OUT)


def register_korean_fonts() -> tuple[str, str]:
    try:
        pdfmetrics.registerFont(UnicodeCIDFont("HYSMyeongJo-Medium"))
        pdfmetrics.registerFont(UnicodeCIDFont("HYGothic-Medium"))
        return "HYSMyeongJo-Medium", "HYGothic-Medium"
    except Exception:
        return "Helvetica", "Helvetica-Bold"


def build_pdf(lines: list[str]) -> None:
    body_font, bold_font = register_korean_fonts()
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="KoreanBody",
            parent=styles["Normal"],
            fontName=body_font,
            fontSize=9,
            leading=13,
            wordWrap="CJK",
        )
    )
    styles.add(
        ParagraphStyle(
            name="KoreanHeading",
            parent=styles["Heading1"],
            fontName=bold_font,
            fontSize=14,
            leading=18,
            spaceBefore=6,
            spaceAfter=6,
            wordWrap="CJK",
        )
    )
    styles.add(
        ParagraphStyle(
            name="KoreanSmall",
            parent=styles["Normal"],
            fontName=body_font,
            fontSize=7,
            leading=10,
            wordWrap="CJK",
        )
    )

    story = []
    in_code = False
    code_lines: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("```"):
            if in_code:
                story.append(Preformatted("\n".join(code_lines), styles["KoreanSmall"]))
                story.append(Spacer(1, 3 * mm))
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        if line.strip().startswith("|"):
            rows, i = parse_table(lines, i)
            if rows:
                table = Table(rows, repeatRows=1)
                table.setStyle(
                    TableStyle(
                        [
                            ("FONT", (0, 0), (-1, -1), body_font, 6),
                            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 2),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                        ]
                    )
                )
                story.append(table)
                story.append(Spacer(1, 3 * mm))
            continue
        text = clean_inline(line)
        if line.startswith("#"):
            story.append(Paragraph(text.lstrip("#").strip(), styles["KoreanHeading"]))
        elif line.startswith("- "):
            story.append(Paragraph("• " + clean_inline(line[2:].strip()), styles["KoreanBody"]))
        elif re.match(r"^\d+\.\s+", line):
            story.append(Paragraph(text, styles["KoreanBody"]))
        else:
            story.append(Paragraph(text, styles["KoreanBody"]))
        story.append(Spacer(1, 1.8 * mm))
        i += 1

    pdf = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    pdf.build(story)


def main() -> None:
    lines = SRC.read_text(encoding="utf-8").splitlines()
    build_docx(lines)
    build_pdf(lines)
    print(f"wrote {DOCX_OUT}")
    print(f"wrote {PDF_OUT}")


if __name__ == "__main__":
    main()
